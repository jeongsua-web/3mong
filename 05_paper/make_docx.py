from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = '/Users/jeongsua/dev/03_fluento/05_paper'

def set_font(run, font_name, size_pt=None, bold=None):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold

def clear_para(para):
    p = para._p
    for r in list(p.findall(qn('w:r'))):
        p.remove(r)

def add_run(para, text, font='휴먼명조', size=None, bold=None):
    run = para.add_run(text)
    set_font(run, font, size, bold)
    return run

def chapter(doc, roman, title):
    p = doc.add_paragraph(style='각 장 제목')
    add_run(p, f'{roman}. {title}', '휴먼고딕', 11)

def sub(doc, text):
    p = doc.add_paragraph(style='바탕글')
    add_run(p, text, '휴먼고딕', 10, True)

def subsub(doc, text):
    p = doc.add_paragraph(style='바탕글')
    add_run(p, text, '휴먼고딕', 10)

def body(doc, text, size=None):
    p = doc.add_paragraph(style='바탕글')
    add_run(p, text, '휴먼명조', size)
    return p

def bullet(doc, text, size=None):
    p = doc.add_paragraph(style='바탕글')
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, '• ' + text, '휴먼명조', size)

def numbered(doc, n, text):
    p = doc.add_paragraph(style='바탕글')
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, f'{n}. ' + text, '휴먼명조')

def code(doc, text):
    p = doc.add_paragraph(style='바탕글')
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Consolas')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')

def fig(doc, fname, caption):
    img_path = os.path.join(BASE_DIR, fname)
    if os.path.exists(img_path):
        p = doc.add_paragraph(style='바탕글')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Inches(3.2))
    cp = doc.add_paragraph(style='바탕글')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cp, caption, '휴먼명조', 9)

def tbl_caption(doc, text):
    p = doc.add_paragraph(style='그림')
    add_run(p, text, '휴먼명조', 9)

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Normal Table'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        set_font(run, '휴먼명조', 9, True)
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(txt)
            set_font(run, '휴먼명조', 9)
    return t

# ── Load template ──────────────────────────────────────────
doc = Document(os.path.join(BASE_DIR, '학술대회용 논문양식(학부생 캡스톤).docx'))

# ── 1. Title table ─────────────────────────────────────────
tt = doc.tables[0]

# Title
tp = tt.rows[0].cells[0].paragraphs[0]
clear_para(tp)
add_run(tp, 'Fluento: AI 기반 적응형 영어 회화 학습 시스템 설계 및 구현', '휴먼고딕', bold=True)

# Authors
ac = tt.rows[2].cells[0]
clear_para(ac.paragraphs[0])
add_run(ac.paragraphs[0], '정수아, 손민경, 김지수', 'Dotum')

clear_para(ac.paragraphs[1])
add_run(ac.paragraphs[1], '신구대학교', 'Dotum', 10)

for p in ac.paragraphs[2:]:
    clear_para(p)

# ── 2. Remove body after title table ──────────────────────
body_el = doc.element.body
elements = list(body_el)
tt_elem = doc.tables[0]._tbl
tt_idx = elements.index(tt_elem)
for elem in elements[tt_idx+1:]:
    if elem.tag != qn('w:sectPr'):
        body_el.remove(elem)

# ── 3. Abstract ────────────────────────────────────────────
doc.add_paragraph('', style='바탕글')
p = doc.add_paragraph(style='바탕글')
add_run(p, '초록', '휴먼고딕', 10, True)
body(doc, '본 논문은 대규모 언어 모델(LLM)을 활용하여 학습자의 영어 수준을 실시간으로 분석하고, 난이도를 자동으로 조절하는 적응형 영어 회화 학습 시스템 Fluento의 설계와 구현을 다룬다. 기존 영어 학습 앱은 정적인 수준 분류와 단방향 콘텐츠 제공 방식에 머물러 학습자의 실력 변화에 즉각적으로 반응하지 못하는 한계가 있다. 본 시스템은 AWS Bedrock의 Claude 3.5 Sonnet 모델을 통해 대화 메시지마다 어휘, 문법, 유창성, 문맥 인식 능력을 분석하고, 규칙 기반 난이도 조절 알고리즘과 결합하여 학습자에게 최적의 대화 경험을 제공한다. Server-Sent Events(SSE)를 이용한 스트리밍 응답, AWS Cognito 기반 인증, Redis를 활용한 점수 캐싱으로 실시간성과 확장성을 동시에 확보하였다. 프론트엔드는 React 19 기반의 직관적 UI로 사용자가 AI 캐릭터와 자연스러운 대화를 나누며 즉각적인 문법 피드백을 받을 수 있도록 설계하였다.')

# ── I. 서론 ───────────────────────────────────────────────
chapter(doc, 'Ⅰ', '서론')
sub(doc, '1.1 연구 배경 및 필요성')
body(doc, '세계화가 가속화됨에 따라 영어 의사소통 능력에 대한 수요는 지속적으로 증가하고 있다. 그러나 전통적인 영어 교육 방식은 강사 의존적이며, 개인의 수준과 학습 속도를 반영하기 어렵다. 기존 상용 영어 학습 앱들은 사전 정의된 커리큘럼에 따라 콘텐츠를 제공하기 때문에 학습자가 이미 알고 있는 내용을 반복하거나 지나치게 어려운 내용에 노출되는 문제가 발생한다.')
body(doc, '대면 학습은 시간과 장소의 제약이 있고, 기존 온라인 학습 플랫폼은 단순한 문제 풀이에 그치는 경우가 많다. 최근 GPT, Claude 등 대규모 언어 모델의 등장으로 자연어 처리 기술이 급격히 발전하였고, 이를 영어 회화 교육에 접목하려는 시도가 증가하고 있으나, 대부분의 연구는 단순 챗봇 수준의 대화 제공에 그치며 학습자의 언어 수준을 동적으로 감지하고 난이도를 실시간으로 조절하는 적응형 학습 시스템은 아직 연구 단계에 머물러 있다.')

sub(doc, '1.2 연구 목적')
body(doc, '본 연구의 목적은 다음과 같다.')
for i, t in enumerate([
    '직관적인 사용자 인터페이스(UI)를 통한 편리한 접근성 제공',
    '사용자 맞춤형 AI 친구 캐릭터 생성 기능 구현',
    '학습자의 메시지를 AI가 분석하여 영어 수준(초급·중급·고급)을 실시간으로 판별',
    '판별된 수준에 따라 AI 응답의 어휘 수준·문장 복잡도·문법 구조를 자동으로 조절',
    '대화 완료 후 문법 오류 및 뉘앙스 개선점을 제공하는 실시간 피드백 기능 구현',
    '보안과 신뢰성을 갖춘 사용자 인증 체계 구축',
], 1):
    numbered(doc, i, t)

sub(doc, '1.3 논문 구성')
body(doc, '본 논문은 2장에서 관련 연구를 검토하고, 3장에서 시스템 아키텍처 및 기술 스택을 설명한다. 4장에서는 프론트엔드 주요 페이지의 기능과 설계를 기술하며, 5장에서는 백엔드 핵심 기능의 구현 방법을 설명한다. 6장에서 보안 및 신뢰성, 7장에서 UI/UX 설계 특징을 다루고, 8장에서 결론 및 향후 연구 방향을 제시한다.')

# ── II. 관련 연구 ─────────────────────────────────────────
chapter(doc, 'Ⅱ', '관련 연구')
sub(doc, '2.1 적응형 학습 시스템 (Adaptive Learning System)')
body(doc, '적응형 학습 시스템은 학습자의 반응을 분석하여 학습 경로나 콘텐츠를 동적으로 조정하는 시스템이다. Knewton, Duolingo 등 상용 서비스는 문항 반응 이론(IRT)이나 지식 추적(Knowledge Tracing) 기법을 활용하나, 자유 형식의 회화 데이터를 실시간으로 분석하는 데는 한계가 있다.')

sub(doc, '2.2 LLM 기반 언어 교육')
body(doc, '대규모 언어 모델을 언어 교육에 활용한 연구는 최근 급증하고 있다. GPT-4를 활용한 에세이 첨삭 연구, 챗봇 기반 회화 연습 시스템 등이 있으나, 다음과 같은 한계가 있다.')
for t in ['학습자 수준 판별이 입력 전 설문에 의존하며 대화 중 동적 갱신이 없음',
          'AI 응답이 학습자 수준에 맞게 자동 조절되지 않음',
          '문법·뉘앙스 피드백이 별도 세션으로 분리되어 학습 흐름이 끊김']:
    bullet(doc, t)
body(doc, '본 연구는 이러한 한계를 극복하기 위해 수준 판별·난이도 조절·피드백을 단일 대화 흐름 안에서 통합 처리한다.')

# ── III. 시스템 설계 ──────────────────────────────────────
chapter(doc, 'Ⅲ', '시스템 설계')
sub(doc, '3.1 전체 시스템 아키텍처')
body(doc, 'Fluento 백엔드는 RESTful API 서버와 SSE 스트리밍 서버를 통합한 단일 Spring Boot 애플리케이션으로 구성된다. 클라이언트는 메시지 전송 후 반환된 스트림 URL로 SSE 연결을 맺어 AI 응답을 실시간으로 수신한다.')
code(doc,
'클라이언트\n'
'│\n'
'├─ POST /api/v1/chat/rooms/{roomId}/messages   (메시지 전송)\n'
'│       └─ messageId + streamUrl 반환\n'
'│\n'
'└─ GET /api/v1/chat/stream?messageId=...       (SSE 구독)\n'
'        ├── ai_start          (스트리밍 시작 알림)\n'
'        ├── level_assessment  (수준 판별 결과)\n'
'        ├── level_adjustment  (난이도 변경 알림, 조건부)\n'
'        ├── ai_chunk × N      (AI 응답 텍스트 스트리밍)\n'
'        └── ai_complete       (완료 + 메시지 ID)\n'
'                └─ [백그라운드] evaluation (문법·뉘앙스 피드백 저장)')

sub(doc, '3.2 기술 스택')
subsub(doc, '프론트엔드')
tbl_caption(doc, '표 1. 프론트엔드 기술 스택')
make_table(doc, ['분류', '기술', '버전'], [
    ['Frontend Framework', 'React', '19.2.4'],
    ['Routing', 'React Router DOM', '7.13.2'],
    ['Build Tool', 'Vite', '8.0.1'],
    ['Module System', 'ES Module', '—'],
    ['State Management', 'React Hooks', '—'],
])
subsub(doc, '백엔드')
tbl_caption(doc, '표 2. 백엔드 기술 스택')
make_table(doc, ['분류', '기술', '버전', '선택 이유'], [
    ['Language', 'Java (JDK)', '17', '안정성, 생태계'],
    ['Framework', 'Spring Boot', '4.0.0', '최신 LTS, 반응형 지원'],
    ['Reactive', 'Spring WebFlux', 'Spring 관리', '비동기 SSE 스트리밍'],
    ['Database', 'PostgreSQL', '17', 'JSONB 지원, 오픈소스'],
    ['Cache', 'Redis', '7', '인메모리 점수 추적'],
    ['ORM', 'Spring Data JPA', 'Spring 관리', '생산성'],
    ['DB 마이그레이션', 'Flyway', 'Spring 관리', '스키마 버전 관리'],
    ['인증', 'AWS Cognito', 'SDK v2 2.30.19', 'OAuth 2.0 / JWT'],
    ['AI', 'AWS Bedrock (Claude 3.5 Sonnet)', 'SDK v2 2.30.19', '고성능 LLM, 스트리밍'],
    ['Build', 'Gradle', '9.0.0', '빠른 빌드, 의존성 관리'],
    ['컨테이너', 'Docker / Docker Compose', '—', '환경 일관성'],
])

sub(doc, '3.3 데이터베이스 스키마')
tbl_caption(doc, '표 3. 데이터베이스 테이블 구조')
make_table(doc, ['테이블', '역할'], [
    ['users', '사용자 계정 (Cognito Sub 매핑)'],
    ['chat_rooms', '채팅방 (현재 수준, 캐릭터 정보)'],
    ['chat_messages', '메시지 (수준, 평가 점수, 피드백 JSONB)'],
    ['level_assessments', '수준 판별 이력 (신뢰도, 지표)'],
])

sub(doc, '3.4 패키지 구조')
code(doc,
'com.fluento\n'
'├── config/          # 보안, Redis, Jackson, Swagger 설정\n'
'├── controller/      # REST 및 SSE 컨트롤러\n'
'├── domain/          # JPA 엔티티 및 Repository\n'
'├── dto/             # 요청·응답 DTO\n'
'├── exception/       # 전역 예외 처리\n'
'└── service/\n'
'    ├── BedrockService          # AWS Bedrock API 호출\n'
'    ├── AIResponseService       # SSE 스트림 전체 흐름 통합\n'
'    ├── LevelAssessmentService  # 수준 판별\n'
'    ├── LevelAdjustmentService  # 난이도 조절 (Redis 기반)\n'
'    └── EvaluationService       # 문법·뉘앙스 평가')

# ── IV. 프론트엔드 ────────────────────────────────────────
chapter(doc, 'Ⅳ', '프론트엔드 주요 페이지 기능 및 설계')

sub(doc, '4.1 로그인 페이지')
subsub(doc, '목적 및 설계')
body(doc, '로그인 페이지는 기존 사용자의 접근을 관리하고, 개인 학습 데이터를 보호하는 첫 번째 보안 관문이다. 사용자는 이메일 주소와 비밀번호를 입력하여 시스템에 접근한다. 화면 상단에는 Fluento 로고가 배치되어 브랜드 아이덴티티를 강조하며, 중앙의 카드 형태 로그인 영역에는 이메일과 비밀번호 두 개의 입력 필드가 있다.')
subsub(doc, '보안 기능')
body(doc, '비밀번호 필드 우측에는 "보기/숨김" 토글 버튼이 위치한다. 로그인 버튼은 보라색(#4E3473)으로 강조되어 주요 행동을 유도하며, 로그인 완료 시 사용자의 이메일 정보가 로컬 저장소에 임시 저장된다. 하단에는 Google과 Apple 소셜 로그인 및 회원가입, 비밀번호 찾기 링크가 제공된다.')
fig(doc, 'login.png', '그림 1. 로그인 화면')

sub(doc, '4.2 회원가입 페이지')
subsub(doc, '목적 및 설계')
body(doc, '회원가입 페이지는 신규 사용자가 안전하게 계정을 생성하는 과정을 관리한다. 사용자 정보의 정확성과 보안을 동시에 보장하도록 단계적 입력 구조로 구성되어 있다.')
subsub(doc, '이메일 인증 시스템')
body(doc, '이메일 주소 입력 후 "인증요청" 버튼을 클릭하면 6자리 인증번호를 발송한다. 인증번호 하단에는 유효 시간(03:00)을 표시하여 사용자에게 시간 제약을 명확히 알리며, 이 타이머는 무차별 대입 공격으로부터 계정을 보호한다. 모든 필드가 올바르게 입력되면 "가입하기" 버튼이 활성화되며, 가입 완료 후 자동으로 로그인 페이지로 이동한다.')
fig(doc, 'signin.png', '그림 2. 회원가입 화면')

sub(doc, '4.3 홈 페이지')
subsub(doc, '목적 및 설계')
body(doc, '홈 페이지는 사용자의 일일 학습 목표를 관리하고 학습 동기를 부여하는 대시보드 역할을 한다. 화면 상단에는 "오늘의 목표에요" 제목과 체크박스 형태의 목표 목록이 나열된다.')
subsub(doc, '목표 추적 기능')
body(doc, '각 목표 옆의 체크박스를 클릭하면 해당 항목이 완료된 것으로 표시된다. 완료된 항목은 회색으로 변하고 텍스트에 취소선이 나타나 시각적 피드백을 제공한다. 페이지 중단부에는 "오늘의 표현 하나 배워보세요" 섹션에 일일 영어 표현과 해석이 제시된다.')
fig(doc, 'home.png', '그림 3. 홈페이지 화면')

sub(doc, '4.4 친구 목록 페이지')
subsub(doc, '목적 및 설계')
body(doc, '친구 목록 페이지는 사용자가 생성한 AI 캐릭터를 관리하고 선택하는 중추 기능이다. "모든 친구"와 "즐겨찾는 친구" 탭과 실시간 검색 입력 창이 제공된다. 각 친구는 카드 형태로 표시되며, 세로 점(⋮) 아이콘으로 즐겨찾기 등록/해제, 삭제 옵션을 제공한다.')
subsub(doc, '모달 창 및 플로팅 버튼')
body(doc, '친구 항목을 클릭하면 프로필 이미지(200×200px)와 메모 입력 영역을 포함한 모달 창이 나타난다. 페이지 우측 하단의 플로팅 "+" 버튼 클릭 시 커스텀 친구 생성 페이지로 이동한다.')
fig(doc, 'friends.png', '그림 4. 친구 목록 페이지와 상세 모달')

sub(doc, '4.5 커스텀 친구 추가 페이지')
subsub(doc, '목적 및 설계')
body(doc, '커스텀 친구 페이지는 사용자가 AI 캐릭터를 직접 생성하는 창작 공간이다. 좌측 프로필 이미지 업로드 영역과 우측 정보 입력 섹션(이름, 성별, 직업, 성격, 메모)으로 구성된다. 이름과 직업이 입력되지 않으면 경고 메시지가 표시되며, 완료 시 자동으로 친구 목록 페이지로 이동한다.')
fig(doc, 'custom.png', '그림 5. 커스텀 친구 생성 화면')

sub(doc, '4.6 채팅 페이지')
subsub(doc, '목적 및 설계')
body(doc, '채팅 페이지는 사용자와 AI 캐릭터 간의 영어 대화를 실현하는 핵심 기능이다. 자연스러운 대화 환경과 함께 실시간 문법 교정을 제공하여 학습 효과를 극대화한다.')
subsub(doc, '실시간 교정 창')
body(doc, '헤더 아래 "실시간 교정 창"이 배치되어 사용자가 입력한 영어 문장의 문법 오류를 즉시 표시한다. 원본 문장은 취소선으로, 수정된 문장은 강조 표시되며 교정 이유 설명이 함께 제공된다. 사용자 메시지는 우측 보라색 말풍선으로, AI 메시지는 좌측 흰색 말풍선으로 표시된다.')
fig(doc, 'chat.png', '그림 6. 채팅 화면과 실시간 교정 창')

# ── V. 백엔드 구현 ────────────────────────────────────────
chapter(doc, 'Ⅴ', '백엔드 구현')

sub(doc, '5.1 AI 수준 판별 (Level Assessment)')
body(doc, '사용자가 메시지를 전송할 때마다 LevelAssessmentService가 해당 메시지를 AWS Bedrock의 Claude 3.5 Sonnet 모델에 전달하여 영어 수준을 판별한다. 분석 항목은 어휘(vocabulary), 문법(grammar), 유창성(fluency), 문맥 인식(contextAwareness)의 4가지이며, 각 항목은 beginner / intermediate / advanced로 분류된다. 프롬프트 설계는 모델이 JSON 형식으로만 응답하도록 지시하여 파싱 오류를 최소화하였다.')
code(doc, '입력 → 현재 수준 + 최근 대화 이력 + 사용자 메시지\n출력 → { detectedLevel, confidence(0~1), indicators: { vocabulary, grammar, fluency, contextAwareness } }')

sub(doc, '5.2 적응형 난이도 조절 (Adaptive Level Adjustment)')
body(doc, 'LevelAdjustmentService는 Redis List(fluento:scores:{roomId})에 저장된 최근 평가 점수를 기반으로 다음 규칙에 따라 난이도를 조절한다.')
tbl_caption(doc, '표 4. 난이도 조절 규칙')
make_table(doc, ['조건', '조치'], [
    ['연속 3개 메시지 점수 ≥ 90', '수준 상향'],
    ['연속 2개 메시지 점수 ≤ 50', '수준 하향'],
    ['그 외', '현재 수준 유지'],
])
body(doc, 'Redis를 사용하여 DB 조회 없이 O(1) 속도로 최근 점수를 조회하고, 수준 변경 시 chat_rooms 테이블의 current_level 컬럼을 갱신한다.')

sub(doc, '5.3 수준별 AI 응답 생성')
body(doc, 'BedrockService는 현재 학습자 수준에 따라 서로 다른 시스템 프롬프트를 AWS Bedrock에 전달한다.')
tbl_caption(doc, '표 5. 수준별 AI 응답 특징')
make_table(doc, ['수준', '시스템 프롬프트 특징'], [
    ['Beginner', '1,000~2,000개 어휘, 5~10 단어 단문, 현재 시제만 사용'],
    ['Intermediate', '3,000~5,000개 어휘, 10~20 단어 문장, 다양한 시제, 관용어 포함'],
    ['Advanced', '5,000개 이상 어휘, 복문·가정법·학문적 표현, 문화적 참조 포함'],
])
body(doc, '응답은 AWS SDK의 ConverseStreamRequest를 통해 스트리밍으로 수신하며, Flux<String> 형태로 리액티브 파이프라인에 통합된다.')

sub(doc, '5.4 실시간 SSE 스트리밍')
body(doc, 'Spring WebFlux의 Flux<ServerSentEvent<>>를 활용하여 비동기 스트리밍을 구현하였다. AIResponseService는 다음 순서로 이벤트를 발행한다.')
for i, t in enumerate(['ai_start — 즉시 발행 (사용자 대기 시간 최소화)',
    'level_assessment — Bedrock 수준 판별 완료 후 발행',
    'level_adjustment — 난이도 변경이 발생한 경우에만 발행',
    'ai_chunk × N — Bedrock 스트리밍 응답 청크를 그대로 전달',
    'ai_complete — 전체 AI 메시지 저장 완료 후 발행'], 1):
    numbered(doc, i, t)
body(doc, '문법·뉘앙스 평가는 ai_complete 발행 이후 별도 스레드(Schedulers.boundedElastic())에서 비동기 실행되므로 SSE 스트림의 응답 시간에 영향을 주지 않는다.')

sub(doc, '5.5 문법·뉘앙스 피드백 (Evaluation)')
body(doc, 'EvaluationService는 대화가 완료된 후 사용자 메시지와 AI 응답을 함께 Bedrock에 전달하여 다음 항목을 평가한다.')
code(doc,
'{\n'
'  "isCorrect": true,\n'
'  "score": 0~100,\n'
'  "feedback": {\n'
'    "grammarIssues": [{ "type", "issue", "correction", "explanation" }],\n'
'    "nuanceIssues":  [{ "type", "context", "suggestion" }],\n'
'    "positivePoints": ["..."],\n'
'    "overallFeedback": "..."\n'
'  }\n'
'}')
body(doc, '평가 결과는 chat_messages 테이블의 JSONB 컬럼에 저장되어, 클라이언트가 메시지 조회 시 피드백을 함께 확인할 수 있다.')

sub(doc, '5.6 인증 및 보안')
body(doc, '운영 환경에서는 AWS Cognito가 발급한 JWT를 Spring Security OAuth2 Resource Server가 검증한다. 개발 환경(dev 프로필)에서는 DevSecurityConfig가 활성화되어 Authorization: Bearer dev-token 헤더로 인증을 우회할 수 있다. 이를 통해 개발 편의성과 운영 보안을 분리하였다.')

sub(doc, '5.7 API 명세')
body(doc, '모든 응답은 통일된 래퍼 포맷을 사용한다.')
code(doc, '{ "success": true,  "data": { ... } }\n{ "success": false, "error": { "code": "CHAT_ROOM_NOT_FOUND", "message": "...", "statusCode": 404 } }')
tbl_caption(doc, '표 6. API 명세')
make_table(doc, ['Method', 'Path', '설명'], [
    ['GET', '/api/v1/health', '헬스체크'],
    ['POST', '/api/v1/auth/register', '회원가입'],
    ['GET', '/api/v1/users/me', '내 정보 조회'],
    ['PUT', '/api/v1/users/profile', '프로필 수정'],
    ['GET', '/api/v1/chat/rooms', '채팅방 목록 조회'],
    ['POST', '/api/v1/chat/rooms', '채팅방 생성'],
    ['POST', '/api/v1/chat/rooms/{roomId}/messages', '메시지 전송'],
    ['GET', '/api/v1/chat/rooms/{roomId}/messages', '메시지 목록 조회'],
    ['GET', '/api/v1/chat/stream?messageId=', 'SSE AI 응답 스트림'],
    ['GET', '/api/v1/chat/rooms/{roomId}/level', '현재 수준 조회'],
])

# ── VI. 보안 및 신뢰성 ───────────────────────────────────
chapter(doc, 'Ⅵ', '보안 및 신뢰성')
sub(doc, '6.1 사용자 인증 체계')
body(doc, '본 플랫폼은 사용자 정보 보호와 신뢰 구축을 최우선으로 한다. 다단계 인증 절차는 다음과 같다.')
for i, t in enumerate(['회원가입: 이메일 인증번호 검증 (6자리, 3분 유효)',
    '로그인: 이메일 + 비밀번호 인증 (JWT 발급)',
    '개인 학습 데이터: localStorage를 통한 세션 관리'], 1):
    numbered(doc, i, t)
body(doc, '비밀번호 보안')
for t in ['입력 필드에서 비밀번호 마스킹 (●●●● 표시)',
          '사용자 선택에 따른 보이기/숨기기 기능',
          '회원가입 시 비밀번호 재확인을 통한 입력 오류 방지']:
    bullet(doc, t)

sub(doc, '6.2 향후 보안 강화 방안')
for t in ['학교 인증 시스템 연동: KUTIS 등 학교 공식 인증 시스템과 연동하여 사용자 신원 자동 확인',
          '보안 질문 도입: 캐릭터 생성 시 보안 질문 설정, 향후 계정 복구 시 본인 인증에 활용',
          '개인정보 보호: 민감한 개인정보는 필요할 때만 표시, 채팅 기록 암호화 저장']:
    bullet(doc, t)

# ── VII. UI/UX 설계 특징 ─────────────────────────────────
chapter(doc, 'Ⅶ', 'UI/UX 설계 특징')
sub(doc, '7.1 색상 및 시각 설계')
for t in ['주색상: #4E3473 (보라색) — 전문성과 신뢰성을 나타내는 브랜드 색상',
          '배경색: #ffffff (흰색) — 깔끔함과 높은 가독성',
          '강조색: 파란색 및 빨간색 — 경고 및 중요 정보 표시',
          '그림자 효과: 부드러운 보라색 그림자로 깊이감 표현',
          '모서리: 둥근 모서리(8–30px)로 친화적 분위기 조성']:
    bullet(doc, t)

sub(doc, '7.2 사용성 설계')
for t in ['직관적 네비게이션: 사이드바로 모든 주요 페이지에 빠르게 접근',
          '조건부 헤더 표시: 친구 목록과 채팅 페이지에서는 헤더를 숨겨 집중력 강화',
          '플로팅 버튼: 주요 행동(친구 추가)을 항상 접근 가능하게 배치',
          '모달 창: 상세 정보를 현재 페이지 맥락 내에서 확인 가능',
          '실시간 피드백: 검색, 필터링, 입력 값 변화에 즉각 반응']:
    bullet(doc, t)

sub(doc, '7.3 반응형 디자인')
for t in ['레이아웃: Flexbox 기반으로 다양한 화면 크기에 적응',
          '접근성: 키보드 네비게이션 지원 및 스크린 리더 호환성 고려']:
    bullet(doc, t)

# ── VIII. 결론 및 향후 연구 ──────────────────────────────
chapter(doc, 'Ⅷ', '결론 및 향후 연구')
sub(doc, '8.1 기대 효과')
body(doc, '본 시스템은 학습자에게 다음과 같은 가치를 제공한다.')
for t in [
    '학습 효율 향상: AI가 매 대화마다 수준을 분석하고 응답 난이도를 자동 조절함으로써 학습자는 Vygotsky의 근접 발달 영역(ZPD) 이론에 부합하는 최적 학습 구간에서 대화를 이어갈 수 있다.',
    '즉각적인 피드백: 대화 직후 문법 오류·뉘앙스 개선점·긍정 포인트를 구조화된 형태로 제공하므로 학습자는 자신의 실수를 빠르게 인지하고 교정할 수 있다.',
    '낮은 심리적 장벽: AI 캐릭터와의 대화이므로 실수에 대한 두려움 없이 자유롭게 표현할 수 있으며, 언제 어디서나 반복 연습이 가능하다.',
    '개인 맞춤형 경험: 초급자와 고급자가 전혀 다른 수준의 대화를 경험하며, 실력이 향상됨에 따라 시스템이 자동으로 더 높은 수준의 콘텐츠를 제공한다.',
]:
    bullet(doc, t)

sub(doc, '8.2 현재 제약사항')
for t in ['백엔드 API 연동 미완료 (데이터베이스 저장 기능 제한)',
          '실제 AI 문법 교정 엔진 미구현 (시뮬레이션 단계)',
          '실시간 채팅 웹소켓 연결 미구현',
          'Flyway 자동 마이그레이션 이슈 (Spring Boot 4 호환 버전 대응 필요)']:
    bullet(doc, t)

sub(doc, '8.3 향후 개발 계획')
subsub(doc, '단기 과제')
for t in ['프론트엔드(iOS/Android) 앱 개발 및 실사용자 테스트',
          'AWS S3 기반 이미지 메시지 전송 구현',
          'WebSocket을 활용한 양방향 통신 구현']:
    bullet(doc, t)
subsub(doc, '중기 과제')
for t in ['수준 판별 정확도 평가 지표 수립 및 A/B 테스트 도입',
          '학습 이력 분석 대시보드 구현 (주간 점수 추이, 오류 유형 통계)',
          '다국어 지원 확장 (일어, 중국어, 스페인어)',
          'RAG(Retrieval-Augmented Generation) 기반 학습 콘텐츠 확장']:
    bullet(doc, t)
subsub(doc, '장기 과제')
for t in ['학습자 행동 데이터를 기반으로 한 파인튜닝 모델 개발',
          '음성 입출력(STT/TTS) 통합으로 발음 교정 기능 추가',
          '그룹 학습 모드: 다수 학습자 간 AI 퍼실리테이터 역할',
          '실시간 수준 판별 모델의 경량화를 통한 온디바이스 추론 연구']:
    bullet(doc, t)

# ── ACKNOWLEDGMENTS ───────────────────────────────────────
p = doc.add_paragraph(style='서론')
add_run(p, 'ACKNOWLEDGMENTS', '휴먼고딕', 10)
p2 = doc.add_paragraph(style='바탕글')
add_run(p2, '본 과제(결과물)은 2026년도 교육부 및 경기도 재원으로 경기 RISE센터의 지원을 받아 수행된 지역혁신중심 대학지원체계(RISE)의 결과입니다.(2026-RISE-09-B21)', '휴먼명조', 9)

# ── REFERENCES ────────────────────────────────────────────
p = doc.add_paragraph(style='바탕글')
add_run(p, 'REFERENCES', '휴먼고딕', 11)
refs = [
    '[1] Gan, W., et al. (2024). Adapting Large Language Models for Education: Foundational Capabilities, Potentials, and Challenges. arXiv preprint arXiv:2401.08664.',
    '[2] Huang, W., Hew, K. F., & Fryer, L. K. (2022). Chatbots for language learning—Are they really useful? Journal of Computer Assisted Learning, 38(1), 237–257.',
    '[3] Wang, Z., et al. (2025). Automatic Proficiency Assessment in L2 English Learners. arXiv preprint arXiv:2505.02615.',
    '[4] Fang, T., et al. (2024). Evaluating Prompting Strategies for Grammatical Error Correction Based on Language Proficiency. arXiv preprint arXiv:2402.15930.',
    '[5] OpenAI. (2023). GPT-4 Technical Report. arXiv preprint arXiv:2303.08774.',
    '[6] Anthropic. (2024). Claude 3.5 Sonnet Model Card Addendum. Anthropic.',
    '[7] AWS. (2024). Amazon Bedrock Developer Guide. Amazon Web Services.',
    '[8] VMware Tanzu. (2024). Spring WebFlux Reference Documentation. Spring Framework.',
    '[9] WHATWG. (2024). Server-Sent Events — HTML Living Standard.',
    '[10] Brown, A., et al. (2023). Interactive Language Learning: The Impact of Real-time Feedback on Student Engagement. Journal of Educational Technology & Society, vol. 24, no. 3, pp. 45–58.',
    '[11] Kim, S., & Park, J. (2022). Conversational AI as a Tool for Language Acquisition. Computer Assisted Language Learning, vol. 35, no. 5, pp. 1–20.',
    '[12] React Documentation. https://react.dev',
    '[13] React Router Documentation. https://reactrouter.com',
    '[14] Vite Documentation. https://vitejs.dev',
]
for r in refs:
    p = doc.add_paragraph(style='바탕글')
    add_run(p, r, '휴먼명조', 9)

# ── Save ──────────────────────────────────────────────────
out = os.path.join(BASE_DIR, 'fluento-paper.docx')
doc.save(out)
print(f'Saved: {out}')
